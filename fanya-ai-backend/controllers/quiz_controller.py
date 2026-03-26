from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import get_jwt_identity, jwt_required
from models.quiz import Quiz, QuizAnswer, QuizAttempt
from models.lesson import Lesson
from models.knowledge import KnowledgeDoc
from extensions import db
import json
import uuid
import re  # 必须引入正则库用于修复 JSON
from concurrent.futures import ThreadPoolExecutor, as_completed  # 引入并发线程池，用于彻底解决截断问题
from datetime import datetime
import pytz
from utils.ai_utils import get_ai_generator
from utils.api_utils import api_response
from utils.rag_utils import loads_rag_chunks, get_rag_context
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

quiz_bp = Blueprint('quiz', __name__)
CNT = pytz.timezone('Asia/Shanghai')

# ==================== 1. 生成测验题接口 (终极方案：并发分批生成 + RAG检索 + JSON 自愈) ====================
@quiz_bp.route('/generate', methods=['POST'])
@jwt_required()
def generate_quiz():
    """并发分批生成20道测验题（彻底解决截断问题，且速度翻倍，支持知识库RAG）"""
    try:
        current_user_id = get_jwt_identity()
        data = request.get_json()
        lesson_id = data.get('lessonId')
        
        if not lesson_id:
            return api_response(code=400, msg='lessonId不能为空')
        
        lesson = Lesson.query.get(lesson_id)
        if not lesson:
            return api_response(code=404, msg='课件不存在')
        
        if lesson.user_id != str(current_user_id):
            return api_response(code=403, msg='无权限操作此课件')
        
        # 获取前端传来的知识库文档ID列表（可选）
        knowledge_doc_ids = data.get('knowledgeDocIds', [])
        
        # 获取讲稿内容
        script_content = ''
        try:
            if lesson.script_content:
                script_data = json.loads(lesson.script_content)
                if isinstance(script_data, list):
                    script_content = '\n'.join([item.get('content', '') for item in script_data])
                else:
                    script_content = str(script_data)
        except:
            script_content = lesson.file_name

        # 加载 RAG chunks，用于精准检索与课件相关内容
        rag_chunks = loads_rag_chunks(lesson.rag_chunks)
        
        # 如果前端传了知识库文档ID，合并这些文档的RAG chunks
        if knowledge_doc_ids:
            for doc_id in knowledge_doc_ids:
                try:
                    kb_doc = KnowledgeDoc.query.get(doc_id)
                    if kb_doc and kb_doc.rag_chunks:
                        kb_chunks = loads_rag_chunks(kb_doc.rag_chunks)
                        rag_chunks.extend(kb_chunks)
                        print(f'[生成测验] 融合知识库文档 [{kb_doc.name}]，新增 {len(kb_chunks)} 个片段')
                except Exception as e:
                    print(f'[生成测验] 加载知识库文档 {doc_id} 失败: {e}')
        
        has_rag = len(rag_chunks) > 0
        print(f'[生成测验] RAG chunks 总数: {len(rag_chunks)}（课件 + 知识库合并）')

        ai_gen = get_ai_generator()

        # JSON 强力自愈修复函数（保留，用于应对偶发的单双引号错误）
        def repair_and_parse_json(raw_text):
            clean_text = re.sub(r'^```(?:json)?\s*|\s*```$', '', raw_text.strip(), flags=re.MULTILINE).strip()
            start_idx = clean_text.find('[')
            end_idx = clean_text.rfind(']')
            if start_idx == -1 or end_idx == -1:
                raise ValueError("未在 AI 响应中找到完整的 [] 数组")
            
            # 移除导致解析崩溃的换行符
            clean_text = clean_text[start_idx:end_idx+1].replace('\n', ' ').replace('\r', '')
            
            try:
                # 尝试一次正常解析
                return json.loads(clean_text)
            except json.JSONDecodeError:
                s = clean_text.replace('\\"', '"')
                s = re.sub(r',\s*\}', '}', s) # 修复尾随逗号
                s = re.sub(r',\s*\]', ']', s) # 修复尾随逗号
                
                # 保护结构引号
                s = re.sub(r'\{\s*"', '{<Q>', s)
                s = re.sub(r'"\s*:', '<Q>:', s)
                s = re.sub(r':\s*"', ':<Q>', s)
                s = re.sub(r'"\s*,', '<Q>,', s)
                s = re.sub(r',\s*"', ',<Q>', s)
                s = re.sub(r'"\s*\}', '<Q>}', s)
                s = re.sub(r'\[\s*"', '[<Q>', s)
                s = re.sub(r'"\s*\]', '<Q>]', s)
                
                # 替换捣乱的内部双引号为单引号
                s = s.replace('"', "'")
                # 恢复结构双引号
                s = s.replace('<Q>', '"')
                
                return json.loads(s)

        # 定义单次生成任务的执行函数
        def generate_batch(task_info):
            q_type = task_info['type']
            q_name = task_info['name']
            count = task_info['count']

            # 用 RAG 检索与该题型最相关的课件内容片段
            if has_rag:
                rag_query = f'{q_name} 知识点 考点 {lesson.file_name}'
                context = get_rag_context(rag_query, rag_chunks, top_k=5)
                print(f'[生成测验] {q_name} RAG 检索到 {len(context)} 字符的相关内容')
            else:
                context = script_content[:2000]
                print(f'[生成测验] {q_name} 无RAG，使用讲稿前2000字')

            # 这个 Prompt 要求极短，只生成 4 道题，绝对不可能触碰最大长度截断！
            prompt = f"""基于以下课件内容，生成 {count} 道【{q_name}】测验题。

课件相关内容：
{context}

请返回一个JSON数组，结构如下：
[
  {{
    "questionType": "{q_type}",
    "questionText": "题目内容",
    "options": ["选项1", "选项2", "选项3", "选项4"],
    "correctAnswer": "答案",
    "explanation": "解析说明",
    "difficulty": "medium"
  }}
]

【各字段填写规范】
- 单选题：options填4个选项如['A.xxx','B.xxx','C.xxx','D.xxx']，correctAnswer填'A'或'B'或'C'或'D'
- 多选题：options填4个选项如['A.xxx','B.xxx','C.xxx','D.xxx']，correctAnswer填如'AB'或'ACD'（多个字母连写）
- 判断题：options必须为['正确','错误']，correctAnswer填'正确'或'错误'
- 计算题：options必须为[]（空数组），correctAnswer填完整的计算结果和步骤摘要，如'结果为120J，计算过程：...'
- 应用题：options必须为[]（空数组），correctAnswer填完整的文字答案，详细说明解题思路和结论

【！！！生死攸关的格式要求！！！】
1. 直接输出以 [ 开头，以 ] 结尾的 JSON 数组。
2. 绝对禁止在题目文本、选项或解析内部使用双引号（"）！如果需要强调请强制使用单引号（'）或中文书名号（《》）。内部双引号会导致系统彻底崩溃！
3. 题目必须严格来源于上方课件内容，不得编造课件中没有的知识点。"""
            
            # 对单个题型的生成，最多允许重试 2 次
            for attempt in range(2):
                try:
                    response = ai_gen.generate_reply(prompt)
                    parsed_data = repair_and_parse_json(response)
                    
                    # 强制校验并修正题型，防止 AI 幻觉乱写 questionType
                    for item in parsed_data:
                        item['questionType'] = q_type
                    
                    print(f"[生成测验] 成功生成 {len(parsed_data)} 道 {q_name}！")
                    return parsed_data
                except Exception as e:
                    print(f"[生成测验] {q_name} 第 {attempt+1} 次失败: {e}")
                    
            # 如果2次都失败了，返回空列表，这不会影响其他类型的题目生成！
            return []

        # 将20道题的巨大任务，拆分为 5 个安全的小任务
        tasks = [
            {"type": "single_choice", "name": "单选题", "count": 4},
            {"type": "multiple_choice", "name": "多选题", "count": 4},
            {"type": "true_false", "name": "判断题", "count": 4},
            {"type": "calculation", "name": "计算题", "count": 4},
            {"type": "application", "name": "应用题", "count": 4}
        ]
        
        print(f'[生成测验] 启动 5 线程并发生成，突破大模型长度限制...')
        all_quiz_data = []
        
        # 使用多线程并发执行 5 个请求！
        with ThreadPoolExecutor(max_workers=5) as executor:
            # 提交所有任务并获取 future 对象
            futures = [executor.submit(generate_batch, task) for task in tasks]
            
            # 等待所有任务完成并收集结果
            for future in as_completed(futures):
                result = future.result()
                if result:
                    all_quiz_data.extend(result)
                
        print(f'[生成测验] 并发生成完毕！共成功收集 {len(all_quiz_data)} 道题')
        
        # 如果所有线程全都失败，才向前端报错
        if len(all_quiz_data) == 0:
            return api_response(code=500, msg='AI生成的格式极度不稳定，请稍后再次尝试')

        # 先删除该课件的旧测验题
        Quiz.query.filter_by(lesson_id=lesson_id).delete()
        db.session.commit()
        print(f'[生成测验] 已清空课件{lesson_id}的旧测验题')
        
        # 保存到数据库
        quiz_objects = []
        print(f'[生成测验] 开始保存到数据库...')
        for idx, item in enumerate(all_quiz_data[:20]): # 确保最多不超过20题
            options_val = item.get('options', [])
            if isinstance(options_val, str):
                try:
                    options_val = json.loads(options_val.replace("'", '"'))
                except:
                    options_val = []
                    
            quiz = Quiz(
                quiz_id=f"quiz_{lesson_id}_{idx}_{uuid.uuid4().hex[:6]}",
                lesson_id=lesson_id,
                question_type=item.get('questionType', 'single_choice'),
                question_text=item.get('questionText', ''),
                options=json.dumps(options_val),
                correct_answer=item.get('correctAnswer', ''),
                explanation=item.get('explanation', ''),
                difficulty=item.get('difficulty', 'medium'),
                points=5,
                sort_order=idx,
                source='ai'
            )
            db.session.add(quiz)
            quiz_objects.append(quiz) # 注意：这里先存对象，不要调 to_dict()
        
        # 核心修复点：先 commit 提交，让数据库自动生成 create_time 等时间戳字段
        db.session.commit()
        print(f'[生成测验] 数据已成功持久化到数据库')
        
        # 提交完成后，时间字段已经有了，此时再转换为字典，绝对不会报错
        quiz_list = [q.to_dict() for q in quiz_objects]
        
        return api_response(data={
            'lessonId': lesson_id,
            'quizCount': len(quiz_list),
            'quizzes': quiz_list
        }, msg='测验题生成成功')
        
    except Exception as e:
        db.session.rollback()
        return api_response(code=500, msg=f'生成失败: {str(e)}')


# ==================== 2. 查询测验题接口 ====================
@quiz_bp.route('/list/<int:lesson_id>', methods=['GET'])
@jwt_required()
def list_quizzes(lesson_id):
    try:
        lesson = Lesson.query.get(lesson_id)
        if not lesson:
            return api_response(code=404, msg='课件不存在')
        quizzes = Quiz.query.filter_by(lesson_id=lesson_id).order_by(Quiz.sort_order).all()
        return api_response(data={
            'lessonId': lesson_id,
            'quizCount': len(quizzes),
            'quizzes': [q.to_dict() for q in quizzes]
        })
    except Exception as e:
        return api_response(code=500, msg=f'查询失败: {str(e)}')


# ==================== 3. 新增题目接口 ====================
@quiz_bp.route('/add', methods=['POST'])
@jwt_required()
def add_quiz():
    try:
        current_user_id = get_jwt_identity()
        data = request.get_json()
        lesson_id = data.get('lessonId')
        lesson = Lesson.query.get(lesson_id)
        if not lesson:
            return api_response(code=404, msg='课件不存在')
        if lesson.user_id != str(current_user_id):
            return api_response(code=403, msg='无权限操作此课件')
        max_order = db.session.query(db.func.max(Quiz.sort_order)).filter_by(lesson_id=lesson_id).scalar() or -1
        quiz = Quiz(
            quiz_id=f"quiz_{lesson_id}_{uuid.uuid4().hex[:8]}",
            lesson_id=lesson_id,
            question_type=data.get('questionType', 'single_choice'),
            question_text=data.get('questionText', ''),
            options=json.dumps(data.get('options', [])),
            correct_answer=data.get('correctAnswer', ''),
            explanation=data.get('explanation', ''),
            difficulty=data.get('difficulty', 'medium'),
            points=data.get('points', 5),
            sort_order=max_order + 1,
            source='manual'
        )
        db.session.add(quiz)
        db.session.commit()
        return api_response(data=quiz.to_dict(), msg='题目添加成功')
    except Exception as e:
        db.session.rollback()
        return api_response(code=500, msg=f'添加失败: {str(e)}')


# ==================== 4. 编辑题目接口 ====================
@quiz_bp.route('/update/<int:quiz_id>', methods=['PUT'])
@jwt_required()
def update_quiz(quiz_id):
    try:
        current_user_id = get_jwt_identity()
        quiz = Quiz.query.get(quiz_id)
        if not quiz:
            return api_response(code=404, msg='题目不存在')
        lesson = Lesson.query.get(quiz.lesson_id)
        if lesson.user_id != str(current_user_id):
            return api_response(code=403, msg='无权限操作此题目')
        data = request.get_json()
        quiz.question_type = data.get('questionType', quiz.question_type)
        quiz.question_text = data.get('questionText', quiz.question_text)
        quiz.options = json.dumps(data.get('options', json.loads(quiz.options or '[]')))
        quiz.correct_answer = data.get('correctAnswer', quiz.correct_answer)
        quiz.explanation = data.get('explanation', quiz.explanation)
        quiz.difficulty = data.get('difficulty', quiz.difficulty)
        quiz.points = data.get('points', quiz.points)
        quiz.update_time = datetime.now(CNT)
        db.session.commit()
        return api_response(data=quiz.to_dict(), msg='题目更新成功')
    except Exception as e:
        db.session.rollback()
        return api_response(code=500, msg=f'更新失败: {str(e)}')


# ==================== 5. 删除题目接口 ====================
@quiz_bp.route('/delete/<int:quiz_id>', methods=['DELETE'])
@jwt_required()
def delete_quiz(quiz_id):
    try:
        current_user_id = get_jwt_identity()
        quiz = Quiz.query.get(quiz_id)
        if not quiz:
            return api_response(code=404, msg='题目不存在')
        lesson = Lesson.query.get(quiz.lesson_id)
        if lesson.user_id != str(current_user_id):
            return api_response(code=403, msg='无权限操作此题目')
        db.session.delete(quiz)
        db.session.commit()
        return api_response(msg='题目删除成功')
    except Exception as e:
        db.session.rollback()
        return api_response(code=500, msg=f'删除失败: {str(e)}')


# ==================== 6. 学生提交答案接口（支持批量提交与高级判分） ====================
@quiz_bp.route('/submit-answer', methods=['POST'])
@jwt_required()
def submit_answer():
    try:
        current_user_id = get_jwt_identity()
        data = request.get_json()
        lesson_id = data.get('lessonId')
        answers = data.get('answers', [])
        
        print(f"[提交答案] 收到请求: lessonId={lesson_id}, answers数量={len(answers)}")
        
        if not lesson_id:
            return api_response(code=400, msg='lessonId不能为空')
        if not answers or not isinstance(answers, list):
            return api_response(code=400, msg='answers必须是数组')
        
        total_score = 0
        correct_count = 0
        wrong_count = 0
        wrong_answers = []
        all_answers = []  # 保存所有题目的答题记录
        
        # 批量处理每道题
        for idx, ans in enumerate(answers):
            quiz_id = ans.get('quizId')
            student_answer = ans.get('answer', '')
            
            # 先尝试用主键查询，如果失败则使用 quiz_id 字段查询
            quiz = Quiz.query.get(quiz_id)
            if not quiz:
                quiz = Quiz.query.filter_by(quiz_id=str(quiz_id)).first()
            if not quiz:
                continue
            
            # 判断答案是否正确
            is_correct = False
            score = 0
            
            # 获取选项列表
            options = json.loads(quiz.options) if quiz.options else []
            
            if quiz.question_type == 'multiple_choice':
                # 多选题：比较选项集合
                # 处理正确答案格式：可能是 "AB" 或 "A,B" 或 "A, B"
                correct_ans = str(quiz.correct_answer).strip().upper()
                # 如果答案中没有逗号，但有多个字母（如"AB"），按单个字母分割
                if ',' not in correct_ans and len(correct_ans) > 1:
                    correct_set = set(correct_ans)
                else:
                    # 按逗号分割并去除空格
                    correct_set = set(a.strip() for a in correct_ans.split(',')) if correct_ans else set()
                
                # 处理学生答案
                if isinstance(student_answer, list):
                    student_set = set(str(a).strip().upper() for a in student_answer)
                else:
                    student_ans = str(student_answer).strip().upper()
                    # 如果答案中没有逗号，但有多个字母（如"AB"），按单个字母分割
                    if ',' not in student_ans and len(student_ans) > 1:
                        student_set = set(student_ans)
                    else:
                        # 按逗号分割并去除空格
                        student_set = set(a.strip() for a in student_ans.split(',')) if student_ans else set()
                
                is_correct = correct_set == student_set
                print(f"[判卷调试] 多选题: 学生答案={student_set}, 正确答案={correct_set}, 结果={is_correct}")
            elif quiz.question_type == 'true_false':
                # 判断题：处理多种答案格式
                student_ans = str(student_answer).strip()
                correct_ans = str(quiz.correct_answer).strip()
                
                def normalize_true_false_answer(ans):
                    ans = ans.upper()
                    if ans in ['正确', '对', '是', 'TRUE', 'YES', 'A']:
                        return 'TRUE'
                    elif ans in ['错误', '错', '否', 'FALSE', 'NO', 'B']:
                        return 'FALSE'
                    return ans
                
                is_correct = normalize_true_false_answer(student_ans) == normalize_true_false_answer(correct_ans)
            elif quiz.question_type == 'calculation':
                # 计算题：待人工/AI深度批改，目前算错或算部分分
                is_correct = False  
                score = 0
            elif quiz.question_type == 'application':
                # 应用题：检查是否有选项
                if options and len(options) > 0:
                    student_ans = str(student_answer).strip().upper()
                    correct_ans = str(quiz.correct_answer).strip().upper()
                    is_correct = student_ans == correct_ans
                else:
                    is_correct = False
                    score = 0
            else:
                # 单选题：支持选项字母和选项文本两种格式的比较
                student_ans = str(student_answer).strip()
                correct_ans = str(quiz.correct_answer).strip()
                
                # 方式1：直接比较（如果correct_answer是字母如'A'）
                if student_ans.upper() == correct_ans.upper():
                    is_correct = True
                # 方式2：将用户答案的字母转换为选项文本再比较
                elif options and len(options) > 0:
                    # 将字母转换为索引（A->0, B->1, C->2, D->3）
                    letter_to_index = {'A': 0, 'B': 1, 'C': 2, 'D': 3}
                    student_index = letter_to_index.get(student_ans.upper())
                    if student_index is not None and student_index < len(options):
                        # 获取用户选择的选项文本
                        student_option_text = options[student_index]
                        # 与正确答案（可能是选项文本）比较
                        is_correct = student_option_text == correct_ans
                        print(f"[判卷调试] 单选题: 用户选择={student_ans}, 索引={student_index}, 选项文本={student_option_text}, 正确答案={correct_ans}, 结果={is_correct}")
            
            if is_correct:
                score = quiz.points
                correct_count += 1
            else:
                wrong_count += 1
                wrong_answers.append({
                    'index': idx,
                    'question': quiz.question_text,
                    'userAnswer': student_answer,
                    'correctAnswer': quiz.correct_answer,
                    'explanation': quiz.explanation
                })
            
            # 记录所有题目的答题情况
            all_answers.append({
                'index': idx,
                'questionType': quiz.question_type,
                'question': quiz.question_text,
                'userAnswer': student_answer,
                'correctAnswer': quiz.correct_answer,
                'isCorrect': is_correct,
                'explanation': quiz.explanation
            })
            
            total_score += score
            
            # 保存答题细粒度记录
            answer_record = QuizAnswer(
                quiz_id=str(quiz_id),
                lesson_id=lesson_id,
                user_id=str(current_user_id),
                student_answer=json.dumps(student_answer) if isinstance(student_answer, list) else str(student_answer),
                is_correct=is_correct,
                score=score
            )
            db.session.add(answer_record)
        
        db.session.commit()
        
        # 计算总分（百分制预估）
        total_questions = len(answers)
        max_score = total_questions * 5  # 假设每题5分
        final_score = int((total_score / max_score) * 100) if max_score > 0 else 0
        
        # AI批改点评（基于分数段生成简单鼓励与建议）
        ai_review = ""
        if final_score >= 90:
            ai_review = "太棒了！你对本节课的内容掌握得非常扎实，所有重点知识都理解到位。继续保持！"
        elif final_score >= 80:
            ai_review = "表现不错！大部分知识点已经掌握，建议针对错题进行复习巩固。"
        elif final_score >= 60:
            ai_review = "基本及格，但还有一些知识点需要加强。建议重新学习相关课件内容，特别是错题涉及的知识点。"
        else:
            ai_review = "需要加油哦！建议重新认真学习本节课件，理解基础概念后再进行测验。"
        
        # 保存本次综合测验尝试记录
        attempt_id = f"attempt_{lesson_id}_{current_user_id}_{int(datetime.now(CNT).timestamp())}"
        attempt = QuizAttempt(
            attempt_id=attempt_id,
            lesson_id=lesson_id,
            user_id=str(current_user_id),
            total_score=final_score,
            correct_count=correct_count,
            wrong_count=wrong_count,
            total_questions=total_questions,
            answers_detail=json.dumps({
                'answers': answers,
                'wrongAnswers': wrong_answers,
                'allAnswers': all_answers,
                'aiReview': ai_review
            })
        )
        db.session.add(attempt)
        db.session.commit()
        
        return api_response(data={
            'score': final_score,
            'correctCount': correct_count,
            'wrongCount': wrong_count,
            'totalScore': total_score,
            'totalQuestions': total_questions,
            'aiReview': ai_review,
            'wrongAnswers': wrong_answers,
            'allAnswers': all_answers,
            'attemptId': attempt_id
        }, msg='答案已提交')
    except Exception as e:
        db.session.rollback()
        return api_response(code=500, msg=f'提交失败: {str(e)}')


# ==================== 7. 获取学生答题记录接口（返回最近一次测验） ====================
@quiz_bp.route('/student-answers/<int:lesson_id>', methods=['GET'])
@jwt_required()
def get_student_answers(lesson_id):
    try:
        current_user_id = get_jwt_identity()
        print(f'[获取答题记录] lesson_id={lesson_id}, user_id={current_user_id}')
        
        # 查询最近一次测验尝试记录
        try:
            latest_attempt = QuizAttempt.query.filter_by(
                lesson_id=lesson_id, 
                user_id=str(current_user_id)
            ).order_by(QuizAttempt.submit_time.desc()).first()
        except Exception as db_error:
            return api_response(code=500, msg=f'数据库错误: {str(db_error)}。请确保已重启后端服务创建表。')
        
        if latest_attempt:
            answers_detail = json.loads(latest_attempt.answers_detail) if latest_attempt.answers_detail else {}
            result = {
                'lessonId': lesson_id,
                'attemptId': latest_attempt.attempt_id,
                'totalScore': latest_attempt.total_score,
                'correctCount': latest_attempt.correct_count,
                'wrongCount': latest_attempt.wrong_count,
                'totalCount': latest_attempt.total_questions,
                'submitTime': latest_attempt.submit_time.strftime('%Y-%m-%d %H:%M:%S'),
                'answersDetail': answers_detail
            }
            return api_response(data=result)
        
        # 如果没有测验记录，返回空
        return api_response(data={
            'lessonId': lesson_id,
            'totalCount': 0,
            'answersDetail': {}
        })
    except Exception as e:
        return api_response(code=500, msg=f'查询失败: {str(e)}')


# ==================== 7.1 获取学生测验历史列表接口 ====================
@quiz_bp.route('/attempts/<int:lesson_id>', methods=['GET'])
@jwt_required()
def get_quiz_attempts(lesson_id):
    """获取学生的所有测验历史记录"""
    try:
        current_user_id = get_jwt_identity()
        attempts = QuizAttempt.query.filter_by(
            lesson_id=lesson_id, 
            user_id=str(current_user_id)
        ).order_by(QuizAttempt.submit_time.desc()).all()
        
        return api_response(data={
            'lessonId': lesson_id,
            'attemptCount': len(attempts),
            'attempts': [a.to_dict() for a in attempts]
        })
    except Exception as e:
        return api_response(code=500, msg=f'查询失败: {str(e)}')


# ==================== 8. 导出Word接口 (已修复综合题适配与鲁棒性) ====================
@quiz_bp.route('/export-word/<int:lesson_id>', methods=['GET'])
@jwt_required()
def export_word(lesson_id):
    """导出测验题为Word文档"""
    try:
        current_user_id = get_jwt_identity()
        lesson = Lesson.query.get(lesson_id)
        if not lesson:
            return api_response(code=404, msg='课件不存在')
        
        quizzes = Quiz.query.filter_by(lesson_id=lesson_id).order_by(Quiz.sort_order).all()
        if not quizzes:
            return api_response(code=404, msg='该课件暂无题目可导出')
        
        doc = Document()
        title = doc.add_heading(f'{lesson.file_name} - 测验题', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 按题型分组
        quiz_by_type = {}
        for quiz in quizzes:
            q_type = quiz.question_type or 'other'
            if q_type not in quiz_by_type:
                quiz_by_type[q_type] = []
            quiz_by_type[q_type].append(quiz)
        
        # 题型映射字典（补充了 comprehensive）
        type_names = {
            'single_choice': '一、单选题',
            'multiple_choice': '二、多选题',
            'true_false': '三、判断题',
            'calculation': '四、计算题',
            'application': '五、应用题',
            'comprehensive': '四、综合题'
        }
        
        # 按照题目顺序循环遍历
        for qtype, qlist in quiz_by_type.items():
            heading_text = type_names.get(qtype, f'其他题型 ({qtype})')
            doc.add_heading(heading_text, level=1)
            
            for quiz in qlist:
                # 题目文本
                q_para = doc.add_paragraph(style='List Number')
                q_para.add_run(quiz.question_text).font.size = Pt(12)
                
                # 选项（安全处理）
                if quiz.options and quiz.options != '[]' and quiz.options != 'null':
                    try:
                        options = json.loads(quiz.options)
                        if isinstance(options, list):
                            for opt in options:
                                doc.add_paragraph(str(opt), style='List Bullet')
                    except Exception:
                        pass
                
                # 答案和解析
                doc.add_paragraph()
                ans_para = doc.add_paragraph()
                ans_para.add_run(f'答案：{quiz.correct_answer or "无"}').bold = True
                
                exp_para = doc.add_paragraph(f'解析：{quiz.explanation or "无"}')
                exp_para.paragraph_format.left_indent = Inches(0.2)
                doc.add_paragraph()
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            attachment_filename=f'{lesson.file_name}_测验题.docx'
        )
        
    except Exception as e:
        print(f"[导出Word] 错误: {str(e)}") # 在后台打印详细错误
        return api_response(code=500, msg=f'导出失败: {str(e)}')


# ==================== 9. 保存智能插旗考点答题记录接口 ====================
@quiz_bp.route('/checkpoint-answer', methods=['POST'])
@jwt_required()
def save_checkpoint_answer():
    """保存学生智能插旗考点的答题记录"""
    try:
        current_user_id = get_jwt_identity()
        data = request.get_json()
        
        lesson_id = data.get('lessonId')
        checkpoint_id = data.get('checkpointId')
        page_num = data.get('pageNum')
        question_text = data.get('questionText', '')
        user_answer = data.get('userAnswer', '')
        correct_answer = data.get('correctAnswer', '')
        is_correct = data.get('isCorrect', False)
        
        print(f"[Checkpoint答题] 接收参数: lesson_id={lesson_id}, checkpoint_id={checkpoint_id}, page_num={page_num}")
        
        if not all([lesson_id, checkpoint_id, page_num is not None]):
            print(f"[Checkpoint答题] 缺少必要参数: lesson_id={lesson_id}, checkpoint_id={checkpoint_id}, page_num={page_num}")
            return api_response(code=400, msg='缺少必要参数')
        
        # 动态导入模型
        from models.progress import CheckpointAnswer
        
        # 检查是否已有记录
        existing = CheckpointAnswer.query.filter_by(
            user_id=str(current_user_id),
            lesson_id=str(lesson_id),
            checkpoint_id=str(checkpoint_id)
        ).first()
        
        if existing:
            # 更新现有记录
            existing.user_answer = user_answer
            existing.correct_answer = correct_answer
            existing.is_correct = is_correct
            existing.question_text = question_text
            existing.submit_time = datetime.now(CNT)
            print(f"[Checkpoint答题] 更新记录: user={current_user_id}, lesson={lesson_id}, checkpoint={checkpoint_id}, is_correct={is_correct}")
        else:
            # 创建新记录
            answer_record = CheckpointAnswer(
                user_id=str(current_user_id),
                lesson_id=str(lesson_id),
                checkpoint_id=str(checkpoint_id),
                page_num=page_num,
                question_text=question_text,
                user_answer=user_answer,
                correct_answer=correct_answer,
                is_correct=is_correct
            )
            db.session.add(answer_record)
            print(f"[Checkpoint答题] 创建记录: user={current_user_id}, lesson={lesson_id}, checkpoint={checkpoint_id}, is_correct={is_correct}")
        
        db.session.commit()
        return api_response(data={'success': True}, msg='答题记录已保存')
        
    except Exception as e:
        db.session.rollback()
        print(f"[Checkpoint答题] 保存失败: {str(e)}")
        return api_response(code=500, msg=f'保存失败: {str(e)}')


# ==================== 10. 获取智能插旗考点答题记录接口 ====================
@quiz_bp.route('/checkpoint-answers/<lesson_id>', methods=['GET'])
@jwt_required()
def get_checkpoint_answers(lesson_id):
    """获取学生在指定课件的所有智能插旗考点答题记录"""
    try:
        current_user_id = get_jwt_identity()
        
        # 动态导入模型
        from models.progress import CheckpointAnswer
        
        answers = CheckpointAnswer.query.filter_by(
            user_id=str(current_user_id),
            lesson_id=str(lesson_id)
        ).order_by(CheckpointAnswer.page_num.asc()).all()
        
        return api_response(data={
            'lessonId': lesson_id,
            'totalCount': len(answers),
            'answers': [a.to_dict() for a in answers]
        })
        
    except Exception as e:
        print(f"[Checkpoint答题] 查询失败: {str(e)}")
        return api_response(code=500, msg=f'查询失败: {str(e)}')